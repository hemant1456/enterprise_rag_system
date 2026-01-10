from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# --- Title Section ---
title = doc.add_heading('Enterprise RAG Test Suite: Project Omniscience', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# --- 1. Complex Technical Documentation ---
doc.add_heading('1. System Architecture', level=1)
doc.add_paragraph(
    "The Omniscience engine operates on a distributed neural-symbolic layer. "
    "Unlike standard architectures, it utilizes the 'Hyperion-7' protocol for "
    "data ingestion, which shards data into 128-bit 'quantum-cells'."
)
doc.add_paragraph(
    "Important: The system's primary latency bottleneck is the 'Sigma-Gate' "
    "located in the US-East-2 region. If latency exceeds 45ms, the system "
    "automatically reverts to the 'Legacy-Fallback' mode."
)

# --- 2. High-Density Tables (Testing Table Parsing) ---
doc.add_heading('2. Regional Performance Metrics', level=1)
table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Region ID'
hdr_cells[1].text = 'Active Nodes'
hdr_cells[2].text = 'Avg Throughput'
hdr_cells[3].text = 'Security Tier'

data = [
    ('EU-WEST-1', '452', '1.2 TB/s', 'Tier 4'),
    ('US-EAST-2', '891', '0.8 TB/s', 'Tier 5'),
    ('AP-SOUTH-1', '120', '2.4 TB/s', 'Tier 3'),
    ('SA-EAST-1', '55', '0.4 TB/s', 'Tier 2')
]

for rid, nodes, tp, sec in data:
    row_cells = table.add_row().cells
    row_cells[0].text = rid
    row_cells[1].text = nodes
    row_cells[2].text = tp
    row_cells[3].text = sec

# --- 3. Nested Policy Information (Testing Multi-hop Reasoning) ---
doc.add_heading('3. Security & Access Protocols', level=1)
doc.add_heading('3.1 Employee Clearances', level=2)
doc.add_paragraph("Clearance levels are determined by the 'Golden-Rule' logic:")
p = doc.add_paragraph("Level Alpha: Full access to Sigma-Gate.", style='List Bullet')
p = doc.add_paragraph("Level Beta: Access to Regional Performance Metrics but NO access to Sigma-Gate.", style='List Bullet')
p = doc.add_paragraph("Level Gamma: Read-only access to log files.", style='List Bullet')

# --- 4. Simulated Log Data (Testing Noise Filtering) ---
doc.add_heading('4. System Event Logs (Last 24 Hours)', level=1)
logs = [
    "2026-01-10 08:00:01 INFO: System Heartbeat OK.",
    "2026-01-10 09:15:22 WARN: Sigma-Gate latency at 42ms.",
    "2026-01-10 10:05:10 ERROR: Unauthorized access attempt by USER_882.",
    "2026-01-10 11:30:00 CRITICAL: US-EAST-2 dropped to 0.2 TB/s. Reverting to Legacy-Fallback.",
    "2026-01-10 14:00:00 INFO: Maintenance complete."
]
log_box = doc.add_paragraph()
log_box.add_run("\n".join(logs)).font.name = 'Courier New'

# --- 5. Potential Hallucination Traps ---
doc.add_heading('5. Appendix: Known Anomalies', level=1)
doc.add_paragraph(
    "Note: While Section 1 states US-East-2 is the bottleneck, internal testing "
    "on Dec 20, 2025, suggested that AP-SOUTH-1 might actually be the bottleneck "
    "during monsoon seasons due to undersea cable vibrations."
)

doc.save('RAG_Enterprise_Dataset.docx')