"""Document processing module for extracting text from DOCX files."""
from docx import Document
from typing import List
import logging

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Processes DOCX documents and extracts text content."""
    
    def __init__(self, file_path: str):
        """
        Initialize the document processor.
        
        Args:
            file_path: Path to the DOCX file
        """
        self.file_path = file_path
    
    def extract_text(self) -> str:
        """
        Extract all text content from the DOCX file.
        
        Returns:
            Extracted text as a single string
            
        Raises:
            FileNotFoundError: If the file doesn't exist
            Exception: If document parsing fails
        """
        try:
            doc = Document(self.file_path)
            text_parts = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            full_text = "\n\n".join(text_parts)
            
            if not full_text.strip():
                logger.warning(f"No text content found in {self.file_path}")
                return ""
            
            logger.info(f"Successfully extracted {len(full_text)} characters from {self.file_path}")
            return full_text
            
        except FileNotFoundError:
            logger.error(f"File not found: {self.file_path}")
            raise
        except Exception as e:
            logger.error(f"Error processing document {self.file_path}: {str(e)}")
            raise
